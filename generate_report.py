#!/usr/bin/env python3
"""Generate the Unlink vs Bermuda SDK engineering briefing as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ============================================================================
# TITLE PAGE
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Unlink SDK vs Bermuda SDK', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Comparative Engineering Briefing')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Date: ').bold = True
meta.add_run('April 7, 2025\n')
meta.add_run('Based on: ').bold = True
meta.add_run('Production integration (Unlink/Whisper) + public repo analysis (Bermuda/BermudaBay)\n')
meta.add_run('Classification: ').bold = True
meta.add_run('Internal Engineering')

doc.add_page_break()

# ============================================================================
# PART 1 — UNLINK SDK STANDALONE BRIEFING
# ============================================================================
doc.add_heading('Part 1 — Unlink SDK Engineering Briefing', level=1)

doc.add_heading('Architecture Overview', level=2)
doc.add_paragraph(
    'The Unlink SDK is a two-tier design:'
)
add_table(
    ['Layer', 'Constructor', 'Purpose'],
    [
        ['High-level', 'createUnlink()', 'Deposit, transfer, withdraw, getBalances, pollTransactionStatus'],
        ['Low-level', 'createUnlinkClient()', 'Raw openapi-fetch client for /transactions/prepare/execute, /transactions/{tx_id}/submit'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'Both layers are needed. The high-level client handles deposit/transfer/withdraw. '
    'The low-level client is required for execute() — the most powerful primitive — because '
    'the high-level client doesn\'t expose it.'
)
doc.add_paragraph(
    'Account model: BIP-39 mnemonic → derives both an EVM account (viem) and an Unlink EdDSA '
    'account (unlinkAccount.fromMnemonic()). The Unlink address format is unlink1…\n\n'
    'Proving model is hybrid: key derivation and EdDSA transaction signing happen client-side, '
    'while ZK proof generation and on-chain settlement are handled by the relayer server.'
)

# -- What Works --
doc.add_heading('What Works ✅', level=2)

doc.add_heading('1. Deposit (sdk.deposit())', level=3)
bullet('ERC-20 → private Unlink pool. Fully functional.')
bullet('sdk.ensureErc20Approval() handles Permit2 allowance automatically.')
bullet('Returns a txId (UUID), which reaches "relayed" status via polling.')
bullet('Needs sufficient gas in the EVM wallet for the on-chain approval + deposit tx.')

doc.add_heading('2. Private Transfer (sdk.transfer())', level=3)
bullet('Single recipient: works reliably. Token + amount + unlink1… address.')
bullet('Batch transfer (multi-recipient, single ZK proof): works via sdk.transfer({ transfers: [...] }).')
bullet('Batch is limited to 2 recipients maximum per ZK proof. For payroll with more recipients, sequential calls are required.', bold_prefix='⚠️ ')
bullet('Sender, recipient, and amount are all shielded — on-chain only the Unlink pool contract is visible.')

doc.add_heading('3. Execute (Low-level API — adapter-based)', level=3)
doc.add_paragraph(
    'The key DeFi integration primitive. This is adapter-based: a specific on-chain contract '
    '(UNLINK_ADAPTER at 0x41BF…) receives funds from the pool, executes arbitrary calldata on '
    'external protocols, and returns outputs back to the pool. The adapter is a trusted intermediary contract.'
)
bullet('Proven working for Uniswap V3 swaps (approve → exactInputSingle) via the Unlink Adapter.')
bullet('Proven working for CCTP V2 bridging (approve → depositForBurn) for cross-chain USDC transfers.')
bullet('Flow: prepare/execute → EdDSA sign message_hash → submit → poll until relayed.')
bullet('Signing uses eddsaSign(spendingPrivateKey, fromDecimal(message_hash)) from the SDK.')

doc.add_heading('4. Withdraw (sdk.withdraw())', level=3)
bullet('Private balance → public EVM address. Works correctly.')
bullet('Opposite of deposit: un-shields tokens.')

doc.add_heading('5. Transaction Polling (sdk.pollTransactionStatus())', level=3)
bullet('Reliable. Terminal states: relayed, processed, failed.')
bullet('Configurable intervalMs and timeoutMs.')
bullet('Typically 10-30 seconds to reach "relayed" on testnet.')

doc.add_heading('6. EdDSA Signing (eddsaSign, fromDecimal)', level=3)
bullet('Works with @zk-kit/eddsa-poseidon under the hood (requires blakejs in the bundle).')
bullet('unlinkAccount.fromMnemonic().getAccountKeys() returns spendingPrivateKey for signing.')
bullet('Vercel gotcha: @zk-kit/eddsa-poseidon/blake-2b must be included via outputFileTracingIncludes in next.config.mjs.', bold_prefix='⚠️ ')

# -- Partial --
doc.add_heading('What Partially Works ⚠️', level=2)

doc.add_heading('7. getBalances (sdk.getBalances())', level=3)
bullet('Returns private pool balances per token.')
bullet('Only reports tokens that were directly deposited or received via transfer. Does NOT report tokens obtained from execute() swaps.', bold_prefix='Critical limitation: ')
bullet('Example: swap USDC → WETH through execute(), WETH is re-deposited to pool, but getBalances() still returns 0 WETH.')
bullet('Workaround: maintain a balance_cache table that tracks last swap outputs.')
bullet('Impact: any DeFi operation creating new token positions has invisible balances.')

doc.add_heading('8. Relayer / Transaction Hashes', level=3)
bullet('The relayer returns UUIDs (e.g. a1b2c3d4-…), not 0x transaction hashes.')
bullet('pollTransactionStatus() result does not reliably include the on-chain 0x hash.')
bullet('Workaround: UUID → link to pool address #internaltx on explorer; 0x → direct tx link.')

# -- Missing --
doc.add_heading('What Doesn\'t Work / Is Missing ❌', level=2)

doc.add_heading('9. Compliance Module', level=3)
bullet('No KYC/KYB integration or compliance API.')
bullet('No transaction screening, sanctions list checking, or travel rule support.')
bullet('No way to flag or restrict addresses at the protocol level.')
bullet('Compliance must be built entirely in the application layer.')

doc.add_heading('10. Real-Time Payments / Streaming', level=3)
bullet('No streaming or continuous payment primitive.')
bullet('All transfers are discrete, one-shot operations.')
bullet('Building streaming would require repeated transfer() calls — not practical at 10-30s latency.')
bullet('No webhooks or event subscriptions — must poll for transaction status.')

doc.add_heading('11. Multi-Chain Native Support', level=3)
bullet('SDK only supports a single environment at a time.')
bullet('Currently only Base Sepolia has a deployed pool.', bold_prefix='Base only: ')
bullet('Cross-chain possible via execute() + CCTP V2 (manual composition, not native).')
bullet('Sender is hidden on source chain; recipient and amount visible on destination.')
bullet('Missing: native cross-chain transfers where both sides are shielded.')

doc.add_heading('12. Burner / HD Addresses', level=3)
bullet('unlinkAccount.fromMnemonic() derives a single deterministic unlink1… address.')
bullet('No HD-wallet-style derivation for multiple Unlink addresses from one mnemonic.', bold_prefix='No HD derivation: ')
bullet('The wallet\'s own key derivation handles the EVM side, but the Unlink address is fixed — one per mnemonic.')
bullet('Multiple addresses require separate mnemonics. No SDK helper for burner/stealth addresses.')

doc.add_heading('13. Adapter Limitations', level=3)
bullet('The Unlink Adapter (0x41BF…) is a single on-chain contract for all execute() calls.')
bullet('Works for simple patterns (approve + swap, approve + bridge).')
bullet('Untested: multi-hop swaps, LP positions, yield vault deposits, flash loans.')
bullet('No documentation on adapter gas limits, reentrancy protections, or supported call patterns.')

doc.add_heading('14. Token Support', level=3)
bullet('Only tokens registered in the pool are supported.')
bullet('On Base Sepolia: USDC (0x036CbD…, 6 decimals) and WETH (0x4200…, 18 decimals).')
bullet('No native ETH support — must use WETH.')
bullet('Adding new tokens requires pool-level registration (not exposed in SDK).')

doc.add_heading('15. Error Handling / Debugging', level=3)
bullet('Error messages often opaque (nested { error: { message } } objects).')
bullet('No transaction simulation or dry-run mode.')
bullet('Failed transactions give minimal context — no revert reasons.')

# -- Contracts table --
doc.add_heading('Key Integration Contracts', level=2)
add_table(
    ['Contract', 'Address', 'Role'],
    [
        ['Unlink Pool', '0x647f9b99af97e4b79DD9Dd6de3b583236352f482', 'ZK privacy pool (Base Sepolia)'],
        ['Unlink Adapter', '0x41BF8f07BC4644055db5BA95c422AAC1Be810Be3', 'Executes external calls from pool'],
        ['USDC', '0x036CbD53842c5426634e7929541eC2318f3dCF7e', '6 decimals'],
        ['WETH', '0x4200000000000000000000000000000000000006', '18 decimals'],
        ['Uniswap V3 Router', '0x94cC0AaC535CCDB3C01d6787D6413C739ae12bc4', 'SwapRouter02 (Base Sepolia)'],
        ['CCTP TokenMessengerV2', '0x8FE6B999Dc680CcFDD5Bf7EB0974218be2542DAA', 'Cross-chain USDC bridge'],
    ]
)

doc.add_page_break()

# ============================================================================
# PART 2 — BERMUDA SDK BRIEFING
# ============================================================================
doc.add_heading('Part 2 — Bermuda SDK Engineering Briefing', level=1)

doc.add_heading('Architecture Overview', level=2)
doc.add_paragraph(
    'Bermuda is a modular privacy layer that sits between wallets and dapps. The core concept is a '
    '"shielded hyperpool" — not just an isolated privacy pool but a full privacy layer providing '
    'P2P shielded transactions, shielded swaps, and seamless access to Ethereum dapps through '
    '"shielding adapters". Bootstrapped by Gnosis.'
)

doc.add_heading('Core Components', level=3)
bullet('Smart contracts (hyperpool, ZK-proof verifiers, adapters for ERC-4626, DEXs, Safe)')
bullet('Client-side prover stack (all ZK proving runs client-side for self-custody)')
bullet('Relayer infrastructure (for tx submission only, not proof generation)')

doc.add_heading('ZK Stack', level=3)
bullet('Noir + Barretenberg (@aztec/bb.js@0.87.6) — primary proving system')
bullet('Circom/snarkjs for browser compatibility via their slimejs library')
bullet('Poseidon2 hashing over BN254 field')
bullet('Schnorr signatures (Noir, Poseidon2-based)')
bullet('BLAKE3 (WASM) and BLAKE2s for key derivation')
bullet('MPT proofs (noir-trie-proofs) — Ethereum state/storage proof verification in Noir')

doc.add_heading('SDK API Surface', level=2)
doc.add_paragraph(
    'The SDK is a private npm package. Based on usage across all public repos, the reconstructed API includes:'
)

bullet('sdk.deposit() — deposit tokens into the privacy pool with ERC-20 Permit support', bold_prefix='Deposit: ')
bullet('sdk.transfer() — shielded P2P transfer with batching via UTXO composition', bold_prefix='Transfer: ')
bullet('sdk.withdraw() — withdraw to public EVM address with auto-unwrap WETH→ETH', bold_prefix='Withdraw: ')
bullet('sdk.findUtxos() — scan on-chain commitment events, decrypt with user keypair', bold_prefix='Balance: ')
bullet('sdk.sumAmounts() — aggregate UTXO values per token', bold_prefix='Sum: ')
bullet('sdk.permit() — gasless ERC-20 approvals', bold_prefix='Permit: ')
bullet('sdk.relay() — submit transactions via relayer', bold_prefix='Relay: ')
bullet('sdk.core.prepareTransact() — generate ZK proofs for 2-in-2-out UTXO transactions', bold_prefix='Core: ')
bullet('sdk.safe.* — Gnosis Safe integration (listTxs, stxHash, shielded proposals)', bold_prefix='Safe: ')
bullet('sdk.KeyPair.fromScalar() — key generation from BLAKE2s-derived scalar', bold_prefix='Keys: ')
bullet('sdk.types.Utxo — UTXO creation, decryption, nullifier generation', bold_prefix='UTXO: ')

doc.add_heading('UTXO Model', level=3)
doc.add_paragraph(
    '2-in-2-out UTXO model. Uses nullifiers to prevent double-spending. Each UTXO contains: '
    'amount, token, keypair, optional safe address, optional type. Commitment events are emitted '
    'on-chain with encrypted outputs, decrypted client-side.'
)

doc.add_heading('Safe (Gnosis) Integration', level=3)
bullet('Full Safe wallet fork with 157 files changed.')
bullet('useBermudaSDK() hook replaces useSafeSDK() as primary SDK interface.')
bullet('Shielded balances displayed alongside regular balances in dashboard.')
bullet('Shield/unshield flows for depositing and withdrawing from Safe.')
bullet('MPECDH (Multi-Party Elliptic Curve Diffie-Hellman) for multi-sig key agreement.')
bullet('Backendless propose() — transaction proposals work without Safe\'s backend.')
bullet('128-byte shielded addresses supported in address validators.')

doc.add_heading('Compliance Roadmap', level=3)
bullet('Proof of Innocence (POI) — optional compliance proofs', bold_prefix='Planned: ')
bullet('Viewing keys — selective disclosure to auditors/regulators without leaking identity', bold_prefix='Planned: ')
bullet('zkKYC — zero-knowledge KYC proofs', bold_prefix='Planned: ')
doc.add_paragraph('No public implementation of these compliance features is visible yet.')

doc.add_heading('DeFi Roadmap', level=3)
doc.add_paragraph('V1: Fully shielded swaps (DEX adapters), ERC-4626 vault adapters, shielded crypto debit card payments.')
doc.add_paragraph('V2: Shielded x402 payments (HTTP-native micropayments), shielded retail checkouts, shielding stack for asset managers, shielded on-off-ramps and remittances.')

doc.add_page_break()

# ============================================================================
# PART 3 — COMPARATIVE ANALYSIS
# ============================================================================
doc.add_heading('Part 3 — Comparative Analysis', level=1)

doc.add_heading('Architecture Comparison', level=2)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Model', 'Account-based (private balances per address)', 'UTXO-based (2-in-2-out, nullifier model)'],
        ['ZK Stack', 'EdDSA-Poseidon (lightweight)', 'Noir + Barretenberg / Circom+snarkjs'],
        ['Proving', 'Hybrid (client signs, server generates ZK proofs & settles)', 'Client-side (all ZK proving in browser/CLI)'],
        ['Smart Account', 'None', 'Full Gnosis Safe fork with shielded multi-sig'],
        ['API Style', 'REST API + openapi-fetch client', 'Direct contract interaction + local proof generation'],
        ['Key Format', 'unlink1… addresses (EdDSA-derived)', '128-byte shielded addresses (spend + encrypt pubkeys)'],
        ['Backed by', 'Unlink team', 'Gnosis (bootstrapped)'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'Key philosophical difference: Unlink uses a hybrid proving model — the client handles key derivation '
    'and EdDSA signing of transaction message hashes, but the relayer generates the ZK proofs and settles '
    'on-chain. This means you trust the relayer for proof integrity but retain custody of signing keys. '
    'Bermuda is a full self-custody privacy layer — all ZK proving happens client-side, with the relayer '
    'only used for transaction submission.'
)

# Feature tables
doc.add_heading('Feature-by-Feature Comparison', level=2)

doc.add_heading('1. Deposit', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Status', '✅ Works', '✅ Works'],
        ['Mechanism', 'sdk.deposit() → relayer settles', 'sdk.deposit() → client proves → relayer submits'],
        ['Approval', 'Permit2 auto-approval', 'ERC-20 Permit (gasless)'],
        ['ETH handling', 'Must use WETH manually', 'Auto-wraps ETH → WETH on deposit'],
        ['Batching', 'Single token per call', 'Can batch approval + deposit via MultiSend'],
    ]
)

doc.add_paragraph()
doc.add_heading('2. Transfer', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Status', '✅ Works', '✅ Works'],
        ['Single', '✅ Fully shielded', '✅ Fully shielded'],
        ['Batch', '⚠️ Limited to 2 recipients per proof', '✅ Multiple via UTXO composition (chainable)'],
        ['Privacy', 'Sender, recipient, amount hidden', 'All hidden + encrypted outputs on-chain'],
    ]
)

doc.add_paragraph()
doc.add_heading('3. Execute / DeFi Integration', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Status', '✅ Works (adapter-based)', '🔜 Roadmap (shielding adapters)'],
        ['Pattern', 'Adapter contract executes calls on behalf of pool', 'Planned: per-protocol shielding adapters'],
        ['Proven', 'Uniswap V3 swaps, CCTP V2 bridging', 'Not yet in production'],
        ['Limitation', 'Single adapter contract; getBalances() misses outputs', 'N/A yet'],
    ]
)

doc.add_paragraph()
doc.add_heading('4. Balance Tracking', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Method', 'Server-side: sdk.getBalances() queries API', 'Client-side: sdk.findUtxos() scans + decrypts events'],
        ['Accuracy', '⚠️ Misses tokens from execute() swaps', '✅ Complete — scans all commitment events'],
        ['Caching', 'None built-in (workaround needed)', 'Disk-based UTXO cache + chain-state crawler'],
        ['Offline sync', 'Not supported', 'GitHub Actions-based event crawler'],
    ]
)

doc.add_paragraph()
doc.add_heading('5. Transaction Hashes', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Return format', '⚠️ UUIDs (not 0x hashes)', '✅ Direct 0x transaction hashes'],
        ['Explorer linking', 'Must fallback to pool address', 'Direct tx links possible'],
    ]
)

doc.add_paragraph()
doc.add_heading('6. Cross-Chain', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Native support', '❌ Base Sepolia only', '❌ Base Sepolia only'],
        ['Manual bridging', '✅ CCTP V2 via execute() (sender hidden)', 'Not attempted'],
        ['Multi-chain roadmap', 'Unknown', 'Architecture is chain-parameterized'],
    ]
)

doc.add_paragraph()
doc.add_heading('7. Compliance', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['KYC/KYB', '❌ Nothing', '🔜 zkKYC (planned)'],
        ['Screening', '❌ Nothing', '🔜 Proof of Innocence (planned)'],
        ['Viewing keys', '❌ Nothing', '🔜 Selective disclosure (planned)'],
        ['Audit trail', '❌ No export', '❌ No export (UTXO cache inspectable)'],
    ]
)

doc.add_paragraph()
doc.add_heading('8. Key Management', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Derivation', 'Single address per mnemonic, no HD derivation', 'Flexible: BLAKE2s scalar → KeyPair, multiple paths'],
        ['Stealth/burner', '❌ Not supported', '⚠️ Possible via different scalars, no helper'],
        ['Peer discovery', 'None', 'sdk.config.peers + .bay aliases'],
        ['ENS integration', '✅ unlink.address text record', '✅ .eth + .bay aliases'],
    ]
)

doc.add_paragraph()
doc.add_heading('9. Smart Account / Multi-sig', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Support', '❌ None — single signer only', '✅ Full Gnosis Safe fork'],
        ['Treasury mgmt', 'Application-layer only', 'Native Safe integration with shielded balances'],
    ]
)

doc.add_paragraph()
doc.add_heading('10. Streaming / Real-Time Payments', level=3)
add_table(
    ['', 'Unlink', 'Bermuda'],
    [
        ['Status', '❌ No streaming primitive', '❌ No streaming primitive'],
        ['Workaround', 'Repeated transfer() (10-30s each)', 'Repeated transfers (similar latency)'],
        ['Roadmap', 'Nothing public', 'Shielded x402 payments (HTTP-native micropayments)'],
    ]
)

doc.add_page_break()

# ============================================================================
# SUMMARY SCORECARD
# ============================================================================
doc.add_heading('Summary Scorecard', level=2)
add_table(
    ['Feature', 'Unlink', 'Bermuda', 'Winner'],
    [
        ['Deposit', '✅', '✅', 'Tie'],
        ['Transfer (single)', '✅', '✅', 'Tie'],
        ['Transfer (batch)', '⚠️ Max 2 recipients', '✅ Chainable UTXOs', 'Bermuda'],
        ['DeFi execute', '✅ Adapter-based, proven', '🔜 Planned', 'Unlink'],
        ['Withdraw', '✅', '✅ (+ auto-unwrap)', 'Bermuda (UX)'],
        ['Balance tracking', '⚠️ Misses execute outputs', '✅ Complete UTXO scan', 'Bermuda'],
        ['Tx hashes', '⚠️ UUIDs', '✅ 0x hashes', 'Bermuda'],
        ['Cross-chain', '⚠️ Manual CCTP, Base only', '❌ Base only, no bridge', 'Unlink (barely)'],
        ['Compliance', '❌', '🔜 POI/zkKYC/viewing keys', 'Bermuda (roadmap)'],
        ['Key management', '❌ No HD, single address', '✅ Flexible derivation', 'Bermuda'],
        ['Smart account', '❌', '✅ Full Safe fork', 'Bermuda'],
        ['Streaming', '❌', '🔜 x402 fork', 'Bermuda (roadmap)'],
        ['Proving model', '⚠️ Hybrid (client signs, server proves)', '✅ Full client-side self-custody', 'Bermuda'],
        ['Production readiness', '✅ Deployed, battle-tested', '⚠️ Testnet, SDK private', 'Unlink'],
        ['Developer experience', '✅ Simple REST API', '⚠️ Complex (UTXO, ZK)', 'Unlink'],
        ['Token support', '2 tokens (USDC, WETH)', '3+ (custom via tokenlist)', 'Bermuda'],
    ]
)

doc.add_paragraph()

# ============================================================================
# BOTTOM LINE
# ============================================================================
doc.add_heading('Bottom Line', level=2)

p = doc.add_paragraph()
r = p.add_run('Unlink')
r.bold = True
p.add_run(
    ' is the pragmatic choice today — simpler API, production-deployed, and the only one with working '
    'DeFi execution (adapter-based swaps and bridges). But it has real gaps: broken balance tracking for '
    'execute outputs, no on-chain tx hashes, batch limited to 2, no multi-sig, no compliance, single '
    'address per mnemonic, and Base Sepolia only.'
)

p = doc.add_paragraph()
r = p.add_run('Bermuda')
r.bold = True
p.add_run(
    ' is architecturally superior — client-side proving (true self-custody), UTXO-based balance tracking '
    'that actually works, Gnosis Safe multi-sig integration, compliance roadmap (POI, zkKYC, viewing keys), '
    'and x402 micropayments. But the SDK is private, DeFi adapters aren\'t shipped yet, and it\'s earlier stage.'
)

doc.add_paragraph()
doc.add_heading('What\'s Needed for Comprehensive DeFi + Cross-Chain + Real-Time Payments', level=2)

needs = [
    ('1. ', 'getBalances() must reflect execute() outputs — biggest gap for DeFi composability.'),
    ('2. ', 'On-chain tx hashes in relayer responses — essential for explorer verification.'),
    ('3. ', 'Multi-environment SDK — native support for pools on multiple chains.'),
    ('4. ', 'Streaming payments primitive — protocol-level token streaming or high-frequency transfer mode.'),
    ('5. ', 'Compliance hooks — address screening, optional KYC gating, audit trail export.'),
    ('6. ', 'Burner/stealth address generation — HD-style derivation of multiple addresses from one mnemonic.'),
    ('7. ', 'Webhooks / event subscriptions — push-based notification instead of polling.'),
    ('8. ', 'Adapter documentation — supported call patterns, gas limits, composition rules.'),
    ('9. ', 'Token registry API — list supported tokens per environment, add custom tokens.'),
    ('10. ', 'Transaction simulation — dry-run execute() calls before committing to ZK proof.'),
]
for prefix, text in needs:
    bullet(text, bold_prefix=prefix)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run(
    'The ideal stack would combine Unlink\'s working execute/adapter pattern with Bermuda\'s client-side '
    'proving, UTXO balance tracking, Safe integration, and compliance primitives — then add streaming '
    '(x402 or Sablier-style) and native multi-chain pools.'
).italic = True

# Save
out_path = os.path.expanduser('~/Desktop/Unlink_vs_Bermuda_SDK_Briefing.docx')
doc.save(out_path)
print(f'✅ Saved to {out_path}')
